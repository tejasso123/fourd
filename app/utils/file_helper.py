import os
from typing import Tuple, List

import aiofiles
import markdown
from phi.model.openai import OpenAIChat
from phi.model.message import Message
import asyncio
import aiofiles.os
from app.core import settings
from docx import Document
import base64

TEMP_FOLDER = './temp_files'
openai_api_key = settings.OPENAI_API_KEY
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'doc', 'docx', 'md', 'jpg', 'jpeg', 'png', 'csv', 'xls', 'xlsx', 'mp3', 'mp4',
                      'mpeg', 'mpga', 'wav', 'webm', 'm4a'}
ALLOWED_MIMETYPES = {'application/pdf', 'text/plain', 'application/msword',
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     'text/markdown', 'image/png', 'image/jpeg'}


async def save_file(file):
    path_name = os.path.join(TEMP_FOLDER, file.filename)
    async with aiofiles.open(path_name, 'wb') as out_file:
        content = await file.read()
        await out_file.write(content)
    print(f"File saved to: {path_name}")
    return path_name


def is_allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


async def preprocess_markdown(file_path):
    async with aiofiles.open(file_path, 'r', encoding='utf-8') as md_file:
        markdown_text = await md_file.read()
    return markdown.markdown(markdown_text)


async def preprocess_text_using_openai(plain_text):
    openai_model = settings.OPENAI_MODEL
    openai_temp = settings.OPENAI_TEMP

    openai_chat = OpenAIChat(
        api_key=openai_api_key, id=openai_model, max_tokens=2000, temperature=openai_temp
    )

    messages = [
        Message(role="system", content="Clean and clarify text for a knowledge base."),
        Message(role="user", content=f"Clean and clarify:\n\n{plain_text}")
    ]

    response = await asyncio.to_thread(openai_chat.invoke, messages=messages)
    return response.choices[0].message.content


def is_valid_mime_type(mime_type: str) -> bool:
    return mime_type.lower() in ALLOWED_MIMETYPES


async def remove_file(path_name):
    if os.path.exists(path_name):
        await aiofiles.os.remove(path_name)


async def extract_docx_text_and_images_async(docx_path: str) -> Tuple[str, List[str]]:
    def extract_sync() -> Tuple[str, List[str]]:
        doc = Document(docx_path)
        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    full_text.append(" | ".join(row_text))

        text = "\n".join(full_text)

        # Extract images
        images = []
        for rel in doc.part._rels.values():
            if "image" in rel.reltype:
                img_data = rel.target_part.blob
                img_base64 = base64.b64encode(img_data).decode("utf-8")
                images.append(f"data:image/png;base64,{img_base64}")

        return text, images

    return await asyncio.to_thread(extract_sync)


def convert_image_to_base64(image_path):
    mime_type = "image/png" if image_path.suffix.lower() == '.png' else "image/jpeg"
    with open(image_path, 'rb') as img_file:
        return f"data:{mime_type};base64,{base64.b64encode(img_file.read()).decode()}"
